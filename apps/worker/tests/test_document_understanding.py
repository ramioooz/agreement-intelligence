from __future__ import annotations

import builtins
import multiprocessing
import platform
import sys
from collections.abc import Mapping, Sequence
from io import BytesIO
from multiprocessing import active_children, get_all_start_methods, get_context
from time import monotonic, sleep
from types import ModuleType
from zipfile import ZIP_DEFLATED, ZipFile

import agreement_intelligence_worker.document_understanding as document_understanding
from agreement_intelligence_worker.document_understanding import DocumentParseError, parse_document
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject
from pytest import MonkeyPatch, raises, skip


def _sleeping_parser_child(**_: object) -> None:
    sleep(0.3)


def _constrained_sleeping_parser_child(**_: object) -> None:
    document_understanding._constrain_child_resources()
    sleep(0.3)


def test_parse_digital_pdf_preserves_page_text_and_stable_citation_anchors() -> None:
    first = parse_document(
        content=_pdf_with_text("Master Agreement\nTermination right"),
        content_type="application/pdf",
        source_checksum="a" * 64,
    )
    second = parse_document(
        content=_pdf_with_text("Master Agreement\nTermination right"),
        content_type="application/pdf",
        source_checksum="a" * 64,
    )

    assert first.diagnostics == ()
    assert first.pages[0].number == 1
    assert [block.text for block in first.pages[0].blocks] == [
        "Master Agreement",
        "Termination right",
    ]
    assert first.pages[0].blocks[0].anchor_id == second.pages[0].blocks[0].anchor_id
    assert first.pages[0].blocks[0].anchor_id.startswith("citation-")


def test_parse_docx_preserves_headings_paragraphs_and_tables() -> None:
    document = Document()
    document.add_heading("Liquidity Provider Agreement", level=1)
    document.add_paragraph("The provider shall make executable prices.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Provider"
    table.rows[0].cells[1].text = "Liquidity Co"
    stream = BytesIO()
    document.save(stream)

    parsed = parse_document(
        content=stream.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        source_checksum="b" * 64,
    )

    assert [(block.kind, block.text) for block in parsed.pages[0].blocks] == [
        ("heading", "Liquidity Provider Agreement"),
        ("paragraph", "The provider shall make executable prices."),
        ("table", "Provider | Liquidity Co"),
    ]


def test_parse_low_text_pdf_reports_ocr_required_without_running_ocr() -> None:
    parsed = parse_document(
        content=_pdf_with_text(""),
        content_type="application/pdf",
        source_checksum="c" * 64,
    )

    assert parsed.diagnostics[0].code == "ocr_required"
    assert parsed.diagnostics[0].page_numbers == (1,)


def test_parse_rejects_a_pdf_with_excessive_pages_before_extracting_text() -> None:
    writer = PdfWriter()
    for _ in range(251):
        writer.add_blank_page(width=612, height=792)
    stream = BytesIO()
    writer.write(stream)

    with raises(ValueError, match="Document parsing was rejected by safety limits."):
        parse_document(
            content=stream.getvalue(),
            content_type="application/pdf",
            source_checksum="d" * 64,
        )


def test_parse_rejects_a_newline_heavy_pdf_before_constructing_excessive_blocks() -> None:
    with raises(ValueError, match="Document parsing was rejected by safety limits."):
        parse_document(
            content=_pdf_with_literal_text("\n".join("x" for _ in range(500_000))),
            content_type="application/pdf",
            source_checksum="h" * 64,
        )


def test_parse_rejects_a_malformed_pdf_without_returning_parser_detail() -> None:
    with raises(ValueError, match="Document parsing was rejected by safety limits."):
        parse_document(
            content=b"%PDF-1.7\nconfidential parser trigger",
            content_type="application/pdf",
            source_checksum="e" * 64,
        )


def test_parse_rechecks_compressed_docx_safety_for_stored_documents() -> None:
    archive = BytesIO()
    with ZipFile(archive, "w", ZIP_DEFLATED) as document:
        document.writestr("[Content_Types].xml", "<Types />")
        document.writestr("_rels/.rels", "<Relationships />")
        document.writestr("word/document.xml", "<w:document />")
        document.writestr("word/media/bomb.bin", b"x" * 1_100_000)

    with raises(DocumentParseError):
        parse_document(
            content=archive.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            source_checksum="f" * 64,
        )


def test_parse_hard_timeout_terminates_the_isolated_parser_child(monkeypatch: MonkeyPatch) -> None:
    if "fork" not in get_all_start_methods():
        skip("fork is unavailable on this platform")
    monkeypatch.setattr(document_understanding, "PARSER_TIMEOUT_SECONDS", 0.01)
    monkeypatch.setattr(
        multiprocessing,
        "get_context",
        lambda _: get_context("fork"),
    )
    monkeypatch.setattr(document_understanding, "_parse_in_child", _sleeping_parser_child)

    child_pids_before = {child.pid for child in active_children()}
    started = monotonic()
    with raises(DocumentParseError):
        parse_document(
            content=_pdf_with_text("timeout fixture"),
            content_type="application/pdf",
            source_checksum="g" * 64,
        )

    assert monotonic() - started < 0.2
    assert {child.pid for child in active_children()} <= child_pids_before


def test_parse_reports_ipc_failure_as_temporarily_unavailable(monkeypatch: MonkeyPatch) -> None:
    class BrokenReceive:
        def poll(self, _: float) -> bool:
            return True

        def recv(self) -> tuple[str, object]:
            raise OSError("temporary IPC failure")

        def close(self) -> None:
            return None

    class Send:
        def close(self) -> None:
            return None

    class Process:
        exitcode = 0

        def start(self) -> None:
            return None

        def is_alive(self) -> bool:
            return False

        def join(self, _: float | None = None) -> None:
            return None

        def close(self) -> None:
            return None

    class Context:
        def Pipe(self, *, duplex: bool) -> tuple[BrokenReceive, Send]:
            assert duplex is False
            return BrokenReceive(), Send()

        def Process(self, **_: object) -> Process:
            return Process()

    monkeypatch.setattr(multiprocessing, "get_context", lambda _: Context())

    with raises(RuntimeError, match="Document parser temporarily unavailable"):
        parse_document(
            content=_pdf_with_text("IPC fixture"),
            content_type="application/pdf",
            source_checksum="i" * 64,
        )


def test_resource_limit_application_failure_reports_parser_unavailability(
    monkeypatch: MonkeyPatch,
) -> None:
    class Send:
        def __init__(self) -> None:
            self.messages: list[tuple[str, object | None]] = []

        def send(self, value: tuple[str, object | None]) -> None:
            self.messages.append(value)

        def close(self) -> None:
            return None

    send = Send()

    class ResourceModule:
        RLIMIT_CPU = 0
        RLIMIT_AS = 1
        RLIM_INFINITY = -1

        @staticmethod
        def getrlimit(_: int) -> tuple[int, int]:
            return (60, 60)

        @staticmethod
        def setrlimit(_: int, __: tuple[int, int]) -> None:
            raise OSError("temporary resource limit failure")

    monkeypatch.setitem(sys.modules, "resource", ResourceModule())

    document_understanding._parse_in_child(
        send=send,
        content=_pdf_with_text("resource fixture"),
        content_type="application/pdf",
        source_checksum="j" * 64,
    )

    assert send.messages == [("unavailable", None)]


def test_missing_resource_module_uses_the_parent_hard_timeout(monkeypatch: MonkeyPatch) -> None:
    if "fork" not in get_all_start_methods():
        skip("fork is unavailable on this platform")
    original_import = builtins.__import__

    def import_without_resource(
        name: str,
        globals: Mapping[str, object] | None = None,
        locals: Mapping[str, object] | None = None,
        fromlist: Sequence[str] | None = (),
        level: int = 0,
    ) -> ModuleType:
        if name == "resource":
            raise ImportError("resource module unavailable")
        return original_import(name, globals, locals, fromlist, level)

    monkeypatch.setattr(builtins, "__import__", import_without_resource)
    monkeypatch.setattr(document_understanding, "PARSER_TIMEOUT_SECONDS", 0.01)
    monkeypatch.setattr(
        multiprocessing,
        "get_context",
        lambda _: get_context("fork"),
    )
    monkeypatch.setattr(
        document_understanding,
        "_parse_in_child",
        _constrained_sleeping_parser_child,
    )

    with raises(DocumentParseError):
        parse_document(
            content=_pdf_with_text("timeout fallback fixture"),
            content_type="application/pdf",
            source_checksum="k" * 64,
        )


def test_macos_uses_the_cpu_limit_and_parent_timeout_without_address_space_limit(
    monkeypatch: MonkeyPatch,
) -> None:
    applied_limits: list[int] = []

    class ResourceModule:
        RLIMIT_CPU = 0
        RLIMIT_AS = 1
        RLIM_INFINITY = -1

        @staticmethod
        def getrlimit(_: int) -> tuple[int, int]:
            return (60, 60)

        @staticmethod
        def setrlimit(limit: int, _: tuple[int, int]) -> None:
            applied_limits.append(limit)

    monkeypatch.setitem(sys.modules, "resource", ResourceModule())
    monkeypatch.setattr(platform, "system", lambda: "Darwin")

    document_understanding._constrain_child_resources()

    assert applied_limits == [ResourceModule.RLIMIT_CPU]


def test_child_memory_error_is_a_permanent_input_rejection(monkeypatch: MonkeyPatch) -> None:
    class Send:
        def __init__(self) -> None:
            self.messages: list[tuple[str, object | None]] = []

        def send(self, value: tuple[str, object | None]) -> None:
            self.messages.append(value)

        def close(self) -> None:
            return None

    def exhaust_memory(**_: object) -> object:
        raise MemoryError

    send = Send()
    monkeypatch.setattr(document_understanding, "_constrain_child_resources", lambda: None)
    monkeypatch.setattr(document_understanding, "_parse_document_untrusted", exhaust_memory)

    document_understanding._parse_in_child(
        send=send,
        content=b"%PDF-1.7\nfixture",
        content_type="application/pdf",
        source_checksum="l" * 64,
    )

    assert send.messages == [("rejected", None)]


def test_child_death_before_a_response_is_a_permanent_input_rejection(
    monkeypatch: MonkeyPatch,
) -> None:
    class Receive:
        def poll(self, _: float) -> bool:
            return False

        def close(self) -> None:
            return None

    class Send:
        def close(self) -> None:
            return None

    class Process:
        exitcode = -9

        def start(self) -> None:
            return None

        def is_alive(self) -> bool:
            return False

        def close(self) -> None:
            return None

    class Context:
        def Pipe(self, *, duplex: bool) -> tuple[Receive, Send]:
            assert duplex is False
            return Receive(), Send()

        def Process(self, **_: object) -> Process:
            return Process()

    monkeypatch.setattr(multiprocessing, "get_context", lambda _: Context())

    with raises(DocumentParseError):
        parse_document(
            content=_pdf_with_text("child death fixture"),
            content_type="application/pdf",
            source_checksum="m" * 64,
        )


def test_clean_child_exit_without_a_response_is_temporarily_unavailable(
    monkeypatch: MonkeyPatch,
) -> None:
    class Receive:
        def poll(self, _: float) -> bool:
            return False

        def close(self) -> None:
            return None

    class Send:
        def close(self) -> None:
            return None

    class Process:
        exitcode = 0

        def start(self) -> None:
            return None

        def is_alive(self) -> bool:
            return False

        def close(self) -> None:
            return None

    class Context:
        def Pipe(self, *, duplex: bool) -> tuple[Receive, Send]:
            assert duplex is False
            return Receive(), Send()

        def Process(self, **_: object) -> Process:
            return Process()

    monkeypatch.setattr(multiprocessing, "get_context", lambda _: Context())

    with raises(RuntimeError, match="Document parser temporarily unavailable"):
        parse_document(
            content=_pdf_with_text("clean child exit fixture"),
            content_type="application/pdf",
            source_checksum="n" * 64,
        )


def _pdf_with_text(text: str) -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_reference = writer._add_object(font)
    resources = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_reference})}
    )
    page[NameObject("/Resources")] = resources
    content = DecodedStreamObject()
    lines = text.split("\n") if text else []
    operators = ["BT", "/F1 12 Tf", "72 720 Td"]
    for index, line in enumerate(lines):
        if index:
            operators.append("0 -18 Td")
        operators.append(f"({line}) Tj")
    operators.append("ET")
    content.set_data("\n".join(operators).encode("ascii"))
    page[NameObject("/Contents")] = writer._add_object(content)
    stream = BytesIO()
    writer.write(stream)
    return stream.getvalue()


def _pdf_with_literal_text(text: str) -> bytes:
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_reference = writer._add_object(font)
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_reference})}
    )
    escaped = (
        text.replace("\\", "\\\\")
        .replace("(", "\\(")
        .replace(")", "\\)")
        .replace("\n", "\\n")
        .replace("\r", "\\r")
    )
    content = DecodedStreamObject()
    content.set_data(f"BT /F1 12 Tf 72 720 Td ({escaped}) Tj ET".encode("ascii"))
    page[NameObject("/Contents")] = writer._add_object(content)
    stream = BytesIO()
    writer.write(stream)
    return stream.getvalue()
