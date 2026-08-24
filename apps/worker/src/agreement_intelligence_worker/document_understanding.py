from __future__ import annotations

import multiprocessing
import platform
import sys
from contextlib import suppress
from dataclasses import dataclass
from hashlib import sha256
from io import BytesIO, StringIO
from time import monotonic
from typing import Any, Final, Literal, cast

from agreement_intelligence_platform.document_safety import (
    MAX_DOCUMENT_COMPRESSED_BYTES,
    DocxArchiveSafetyError,
    validate_docx_archive,
)
from docx import Document
from pypdf import PdfReader

BlockKind = Literal["heading", "paragraph", "list_item", "table"]
EvidenceTrust = Literal["untrusted"]

MAX_PDF_PAGES: Final = 250
MAX_PDF_OBJECTS: Final = 20_000
MAX_PDF_BLOCKS: Final = 10_000
MAX_PDF_ROOT_OBJECT_RECOVERY: Final = 1_000
MAX_PDF_RECURSION_DEPTH: Final = 1_000
MAX_EXTRACTED_CHARACTERS: Final = 1_000_000
MAX_DOCX_BLOCKS: Final = 10_000
MAX_DOCX_TABLE_ROWS: Final = 2_000
MAX_DOCX_TABLE_CELLS: Final = 20_000
PARSER_TIMEOUT_SECONDS: Final = 10.0
PARSER_CPU_SECONDS: Final = 12
PARSER_ADDRESS_SPACE_BYTES: Final = 512 * 1024 * 1024
_PROCESS_TERMINATION_GRACE_SECONDS: Final = 0.2


class DocumentParseError(ValueError):
    """An untrusted document exceeded parser policy or could not be parsed safely."""

    category = "document_parse_rejected"

    def __init__(self) -> None:
        super().__init__("Document parsing was rejected by safety limits.")


class DocumentParserUnavailableError(RuntimeError):
    """The parser process could not run and the job can be retried."""

    def __init__(self) -> None:
        super().__init__("Document parser temporarily unavailable.")


@dataclass(frozen=True)
class CitationAnchor:
    anchor_id: str
    source_checksum: str
    page_number: int
    block_index: int
    start_offset: int
    end_offset: int


@dataclass(frozen=True)
class DocumentBlock:
    anchor_id: str
    kind: BlockKind
    text: str
    start_offset: int
    end_offset: int
    trust: EvidenceTrust = "untrusted"


@dataclass(frozen=True)
class DocumentPage:
    number: int
    blocks: tuple[DocumentBlock, ...]


@dataclass(frozen=True)
class DocumentDiagnostic:
    code: str
    message: str
    page_numbers: tuple[int, ...]


@dataclass(frozen=True)
class ParsedDocument:
    source_checksum: str
    pages: tuple[DocumentPage, ...]
    diagnostics: tuple[DocumentDiagnostic, ...]

    @property
    def citations(self) -> tuple[CitationAnchor, ...]:
        return tuple(
            CitationAnchor(
                anchor_id=block.anchor_id,
                source_checksum=self.source_checksum,
                page_number=page.number,
                block_index=index,
                start_offset=block.start_offset,
                end_offset=block.end_offset,
            )
            for page in self.pages
            for index, block in enumerate(page.blocks)
        )


def parse_document(*, content: bytes, content_type: str, source_checksum: str) -> ParsedDocument:
    if len(content) > MAX_DOCUMENT_COMPRESSED_BYTES:
        raise DocumentParseError()
    return _parse_in_isolated_process(
        content=content,
        content_type=content_type,
        source_checksum=source_checksum,
    )


def _parse_in_isolated_process(
    *, content: bytes, content_type: str, source_checksum: str
) -> ParsedDocument:
    context = multiprocessing.get_context("spawn")
    receive, send = context.Pipe(duplex=False)
    process = context.Process(
        target=_parse_in_child,
        kwargs={
            "send": send,
            "content": content,
            "content_type": content_type,
            "source_checksum": source_checksum,
        },
    )
    try:
        try:
            process.start()
        except (OSError, ValueError):
            raise DocumentParserUnavailableError() from None
        send.close()
        deadline = monotonic() + PARSER_TIMEOUT_SECONDS
        while monotonic() < deadline:
            remaining = deadline - monotonic()
            if receive.poll(max(0.0, min(remaining, 0.1))):
                status, payload = cast(tuple[str, object], receive.recv())
                process.join(_PROCESS_TERMINATION_GRACE_SECONDS)
                if process.is_alive():
                    _stop_parser_process(process)
                if status == "ok" and isinstance(payload, ParsedDocument):
                    return payload
                if status == "rejected":
                    raise DocumentParseError()
                raise DocumentParserUnavailableError()
            if not process.is_alive():
                if process.exitcode == 0:
                    raise DocumentParserUnavailableError()
                raise DocumentParseError()
        _stop_parser_process(process)
        raise DocumentParseError()
    except (EOFError, OSError):
        raise DocumentParserUnavailableError() from None
    finally:
        send.close()
        receive.close()
        if process.is_alive():
            _stop_parser_process(process)
        if process.exitcode is not None:
            process.close()


def _parse_in_child(*, send: Any, content: bytes, content_type: str, source_checksum: str) -> None:
    try:
        _constrain_child_resources()
    except Exception:
        _send_child_result(send, "unavailable", None)
        return
    try:
        parsed = _parse_document_untrusted(
            content=content,
            content_type=content_type,
            source_checksum=source_checksum,
        )
    except MemoryError:
        _send_child_result(send, "rejected", None)
    except OSError:
        _send_child_result(send, "unavailable", None)
    except Exception:
        _send_child_result(send, "rejected", None)
    else:
        _send_child_result(send, "ok", parsed)
    finally:
        send.close()


def _send_child_result(send: Any, status: str, payload: object | None) -> None:
    with suppress(BrokenPipeError, OSError):
        send.send((status, payload))


def _constrain_child_resources() -> None:
    sys.setrecursionlimit(MAX_PDF_RECURSION_DEPTH)
    try:
        import resource
    except ImportError:
        return
    _set_resource_limit(resource, resource.RLIMIT_CPU, PARSER_CPU_SECONDS)
    if platform.system() == "Darwin":
        return
    _set_resource_limit(resource, resource.RLIMIT_AS, PARSER_ADDRESS_SPACE_BYTES)


def _set_resource_limit(resource_module: Any, limit: int, maximum: int) -> None:
    _, current_hard = resource_module.getrlimit(limit)
    hard = maximum if current_hard == resource_module.RLIM_INFINITY else min(maximum, current_hard)
    resource_module.setrlimit(limit, (min(maximum, hard), hard))


def _stop_parser_process(process: Any) -> None:
    if not process.is_alive():
        process.join()
        return
    process.terminate()
    process.join(_PROCESS_TERMINATION_GRACE_SECONDS)
    if process.is_alive():
        kill = getattr(process, "kill", None)
        if callable(kill):
            kill()
        process.join(_PROCESS_TERMINATION_GRACE_SECONDS)


def _parse_document_untrusted(
    *, content: bytes, content_type: str, source_checksum: str
) -> ParsedDocument:
    if content_type == "application/pdf":
        pages = _parse_pdf(content, source_checksum)
        diagnostics = _pdf_diagnostics(pages)
    elif content_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        pages = (_parse_docx(content, source_checksum),)
        diagnostics = ()
    else:
        raise ValueError(f"Unsupported document content type: {content_type}")
    return ParsedDocument(
        source_checksum=source_checksum,
        pages=pages,
        diagnostics=diagnostics,
    )


def _parse_pdf(content: bytes, source_checksum: str) -> tuple[DocumentPage, ...]:
    reader = PdfReader(
        BytesIO(content),
        strict=True,
        root_object_recovery_limit=MAX_PDF_ROOT_OBJECT_RECOVERY,
    )
    object_count = sum(len(objects) for objects in reader.xref.values()) + sum(
        len(objects) for objects in reader.xref_objStm.values()
    )
    if object_count > MAX_PDF_OBJECTS or len(reader.pages) > MAX_PDF_PAGES:
        raise DocumentParseError()
    pages: list[DocumentPage] = []
    extracted_characters = 0
    block_count = 0
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        extracted_characters += len(text)
        if extracted_characters > MAX_EXTRACTED_CHARACTERS:
            raise DocumentParseError()
        lines = _bounded_pdf_lines(text, remaining_blocks=MAX_PDF_BLOCKS - block_count)
        block_count += len(lines)
        pages.append(_page_from_text(page_number, lines, source_checksum))
    return tuple(pages)


def _bounded_pdf_lines(text: str, *, remaining_blocks: int) -> list[str]:
    lines: list[str] = []
    for raw_line in StringIO(text):
        line = raw_line.strip()
        if not line:
            continue
        if len(lines) >= remaining_blocks:
            raise DocumentParseError()
        lines.append(line)
    return lines


def _parse_docx(content: bytes, source_checksum: str) -> DocumentPage:
    try:
        validate_docx_archive(content)
    except DocxArchiveSafetyError as error:
        raise DocumentParseError() from error
    document = Document(BytesIO(content))
    items: list[tuple[BlockKind, str]] = []
    extracted_characters = 0
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = paragraph.style.name.lower() if paragraph.style is not None else ""
        if style_name.startswith("heading"):
            kind: BlockKind = "heading"
        elif style_name.startswith("list"):
            kind = "list_item"
        else:
            kind = "paragraph"
        extracted_characters = _append_docx_item(items, kind, text, extracted_characters)
    table_rows = 0
    table_cells = 0
    for table in document.tables:
        rows: list[str] = []
        for row in table.rows:
            table_rows += 1
            table_cells += len(row.cells)
            if table_rows > MAX_DOCX_TABLE_ROWS or table_cells > MAX_DOCX_TABLE_CELLS:
                raise DocumentParseError()
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        table_text = "\n".join(row for row in rows if row.strip())
        if table_text:
            extracted_characters = _append_docx_item(
                items, "table", table_text, extracted_characters
            )
    return _page_from_items(1, items, source_checksum)


def _append_docx_item(
    items: list[tuple[BlockKind, str]], kind: BlockKind, text: str, extracted_characters: int
) -> int:
    if len(items) >= MAX_DOCX_BLOCKS or extracted_characters + len(text) > MAX_EXTRACTED_CHARACTERS:
        raise DocumentParseError()
    items.append((kind, text))
    return extracted_characters + len(text)


def _page_from_text(page_number: int, lines: list[str], source_checksum: str) -> DocumentPage:
    return _page_from_items(
        page_number,
        [
            ("heading" if index == 0 and _looks_like_heading(line) else "paragraph", line)
            for index, line in enumerate(lines)
        ],
        source_checksum,
    )


def _page_from_items(
    page_number: int, items: list[tuple[BlockKind, str]], source_checksum: str
) -> DocumentPage:
    offset = 0
    blocks: list[DocumentBlock] = []
    for block_index, (kind, text) in enumerate(items):
        start_offset = offset
        end_offset = start_offset + len(text)
        blocks.append(
            DocumentBlock(
                anchor_id=_anchor_id(
                    source_checksum,
                    page_number,
                    block_index,
                    start_offset,
                    end_offset,
                ),
                kind=kind,
                text=text,
                start_offset=start_offset,
                end_offset=end_offset,
            )
        )
        offset = end_offset + 1
    return DocumentPage(number=page_number, blocks=tuple(blocks))


def _anchor_id(
    source_checksum: str,
    page_number: int,
    block_index: int,
    start_offset: int,
    end_offset: int,
) -> str:
    value = ":".join(
        (source_checksum, str(page_number), str(block_index), str(start_offset), str(end_offset))
    )
    return f"citation-{sha256(value.encode()).hexdigest()[:24]}"


def _looks_like_heading(text: str) -> bool:
    return len(text) <= 120 and not text.endswith((".", ";", ":"))


def _pdf_diagnostics(pages: tuple[DocumentPage, ...]) -> tuple[DocumentDiagnostic, ...]:
    low_text_pages = tuple(page.number for page in pages if not page.blocks)
    if not low_text_pages:
        return ()
    return (
        DocumentDiagnostic(
            code="ocr_required",
            message="This PDF has insufficient embedded text and requires OCR.",
            page_numbers=low_text_pages,
        ),
    )
